import streamlit as st
import docx
from docx.shared import Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import pdfplumber
import io
import re
import copy

st.set_page_config(page_title="材料驗收單自動生成器", layout="centered")

st.title("📋 材料驗收單自動生成器")

col1, col2 = st.columns(2)
with col1:
    class_name = st.text_input("班級名稱", value="115W0149-飲調輕食暨餐飲創業(桃園)-第1期")
with col2:
    date_str = st.text_input("進貨日期", value="115/02/03")

st.markdown("---")

uploaded_pdf = st.file_uploader("1. 上傳材料明細 PDF (.pdf)", type=["pdf"])
uploaded_doc = st.file_uploader("2. 上傳空白 Word 範本 (.docx)", type=["docx"])
uploaded_images = st.file_uploader(
    "3. 上傳材料照片 (檔名請含數字，如 1.jpg 或 (1).jpg)", 
    type=["jpg", "jpeg", "png"], 
    accept_multiple_files=True
)

def parse_pdf_items(pdf_file):
    """精準解析 PDF 表格與數量"""
    items = {}
    with pdfplumber.open(pdf_file) as pdf:
        for page in pdf.pages:
            tables = page.extract_tables()
            for table in tables:
                for row in table:
                    if not row: continue
                    clean_row = [str(c).replace('\n', ' ').strip() if c else "" for c in row]
                    if clean_row[0].isdigit():
                        p_num = clean_row[0]
                        p_name = clean_row[1] if len(clean_row) > 1 else ""
                        p_spec = clean_row[2] if len(clean_row) > 2 else ""
                        
                        p_qty = "1"
                        if len(clean_row) >= 6:
                            p_qty = clean_row[4]
                        elif len(clean_row) == 5:
                            p_qty = clean_row[3]
                        
                        items[p_num] = {
                            "name": p_name,
                            "spec": p_spec,
                            "qty": p_qty
                        }
            
            if not items:
                text = page.extract_text()
                if text:
                    for line in text.split('\n'):
                        parts = line.split()
                        if parts and parts[0].isdigit():
                            p_num = parts[0]
                            p_name = parts[1] if len(parts) > 1 else ""
                            p_spec = parts[2] if len(parts) > 2 else ""
                            p_qty = parts[4] if len(parts) > 5 else (parts[3] if len(parts) > 4 else "1")
                            items[p_num] = {
                                "name": p_name,
                                "spec": p_spec,
                                "qty": p_qty
                            }
    return items

def duplicate_last_two_rows(table):
    """深層複製表格最後兩列 (照片列與文字列)，達成無縫動態新增"""
    # 複製倒數第二列 (照片列)
    row_img = table.rows[-2]
    new_row_img = table.add_row()
    for idx, cell in enumerate(row_img.cells):
        new_row_img.cells[idx].text = cell.text
        
    # 複製最後一列 (文字列)
    row_txt = table.rows[-2] # 剛加完一列後的倒數第二列
    new_row_txt = table.add_row()
    for idx, cell in enumerate(row_txt.cells):
        new_row_txt.cells[idx].text = cell.text

if st.button("🚀 產生驗收單文件"):
    if not uploaded_doc or not uploaded_pdf:
        st.error("請確保上傳了 Word 範本與 PDF 明細檔！")
    else:
        # A. 解析 PDF
        items_dict = parse_pdf_items(uploaded_pdf)
        total_items = len(items_dict) if items_dict else 0
        
        # B. 建立照片對應表
        image_map = {}
        if uploaded_images:
            for img in uploaded_images:
                match = re.search(r'\((\d+)\)', img.name)
                if match:
                    p_num = match.group(1)
                else:
                    nums = re.findall(r'\d+', img.name)
                    p_num = nums[-1] if nums else None
                
                if p_num:
                    image_map[p_num] = img

        # C. 載入 Word 範本
        doc = docx.Document(uploaded_doc)
        table = doc.tables[0] # 取得主要表格

        # D. 動態檢查：若 PDF 數量超過 6 個，自動增加對應的表格列數
        # 範本預設容納 6 個品號 (每 2 個品號占 2 列)
        if total_items > 6:
            extra_items = total_items - 6
            # 每多 1~2 個品號，就新增一組 (照片列+文字列)
            needed_pairs = (extra_items + 1) // 2 
            for _ in range(needed_pairs):
                duplicate_last_two_rows(table)

        # E. 重新重構表格內容：填入抬頭與全套品號 (包含超出的品號)
        # 先填寫頂部抬頭資訊
        table.rows[0].cells[0].text = class_name
        table.rows[0].cells[1].text = f"進貨日：{date_str}"

        # 遍歷填寫所有品號與照片
        # 表格結構：Row 1 是照片(1,2), Row 2 是文字(1,2) ; Row 3 是照片(3,4), Row 4 是文字(3,4) ...
        current_p_num = 1
        row_idx = 1
        
        while current_p_num <= max(total_items, 6) and row_idx < len(table.rows):
            # 第一欄 (左側)
            p1_str = str(current_p_num)
            # 填寫左側照片
            cell_img1 = table.rows[row_idx].cells[0]
            cell_img1.text = ""
            if p1_str in image_map:
                p = cell_img1.paragraphs[0]
                p.add_run().add_picture(image_map[p1_str], width=Inches(2.2))
            
            # 填寫左側文字
            cell_txt1 = table.rows[row_idx + 1].cells[0]
            if p1_str in items_dict:
                info = items_dict[p1_str]
                cell_txt1.text = f"品號{p1_str} {info['name']} {info['spec']} 數量:{info['qty']}"
            else:
                cell_txt1.text = f"品號{p1_str} 品名 規格 數量"

            # 第二欄 (右側)
            p2_str = str(current_p_num + 1)
            cell_img2 = table.rows[row_idx].cells[1]
            cell_img2.text = ""
            if p2_str in image_map:
                p = cell_img2.paragraphs[0]
                p.add_run().add_picture(image_map[p2_str], width=Inches(2.2))
            
            cell_txt2 = table.rows[row_idx + 1].cells[1]
            if p2_str in items_dict:
                info = items_dict[p2_str]
                cell_txt2.text = f"品號{p2_str} {info['name']} {info['spec']} 數量:{info['qty']}"
            else:
                cell_txt2.text = f"品號{p2_str} 品名 規格 數量"

            # 指針前進 2 個品號、跳過 2 列
            current_p_num += 2
            row_idx += 2

        # F. 加會驗結果
        p = doc.add_paragraph()
        p.add_run("會驗結果：經確認，到貨物品均於效期內，且與規格相符。").bold = True

        # G. 輸出下載
        bio = io.BytesIO()
        doc.save(bio)
        
        st.success(f"🎉 驗收單文件生成成功！共處理 {total_items} 個品號，表格已自動動態調整！")
        st.download_button(
            label="📥 下載完成的 Word 檔",
            data=bio.getvalue(),
            file_name=f"驗收紀錄_{class_name}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
